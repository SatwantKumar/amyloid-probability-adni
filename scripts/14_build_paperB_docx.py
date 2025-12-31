from __future__ import annotations

import argparse
import subprocess
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


_PAGE_BREAK_OPENXML = """```{=openxml}
<w:p><w:r><w:br w:type="page"/></w:r></w:p>
```"""


def _add_page_numbers(doc: Document) -> None:
    section = doc.sections[0]
    footer = section.footer
    if footer.paragraphs:
        p = footer.paragraphs[0]
        p.clear()
    else:
        p = footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "

    fld_char_sep = OxmlElement("w:fldChar")
    fld_char_sep.set(qn("w:fldCharType"), "separate")

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_sep)
    run._r.append(fld_char_end)


def _set_document_style(
    doc: Document,
    *,
    font_name: str,
    font_size_pt: int,
    line_spacing: WD_LINE_SPACING,
) -> None:
    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(font_size_pt)
    style.paragraph_format.line_spacing_rule = line_spacing
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)

    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        if style_name in doc.styles:
            s = doc.styles[style_name]
            s.font.name = font_name
            s.paragraph_format.space_before = Pt(6)
            s.paragraph_format.space_after = Pt(0)


def _clean_main_manuscript_markdown(text: str) -> str:
    lines = text.splitlines()
    if lines and lines[0].startswith("# Paper B — Manuscript draft"):
        lines = lines[1:]
        while lines and not lines[0].strip():
            lines = lines[1:]
    return "\n".join(lines).strip() + "\n"


def _extract_main_figure_legends(fig_legends_md: str) -> str:
    lines = fig_legends_md.splitlines()

    start = None
    end = None
    for i, line in enumerate(lines):
        if line.startswith("## Figure 1."):
            start = i
            break
    if start is None:
        raise ValueError("Could not find '## Figure 1.' in figure legends markdown.")

    for i, line in enumerate(lines):
        if line.startswith("## eFigure 1."):
            end = i
            break
    if end is None:
        end = len(lines)

    main = "\n".join(lines[start:end]).strip()

    # Demote headings so they sit under a 'Figure Legends' section.
    main = main.replace("\n## ", "\n### ")
    if main.startswith("## "):
        main = "### " + main[3:]
    return main + "\n"


def _demote_table_heading(table_md: str) -> str:
    lines = table_md.splitlines()
    if lines and lines[0].startswith("# "):
        lines[0] = "### " + lines[0][2:].strip()
    return "\n".join(lines).strip() + "\n"


def _build_main_markdown(
    *,
    manuscript_md_path: Path,
    figure_legends_md_path: Path,
    table1_md_path: Path,
) -> str:
    manuscript = _clean_main_manuscript_markdown(manuscript_md_path.read_text(encoding="utf-8"))

    # Page breaks before core front-matter sections.
    for heading in ["## Key Points", "## Abstract", "## Introduction", "## References"]:
        manuscript = manuscript.replace(heading, f"{_PAGE_BREAK_OPENXML}\n\n{heading}", 1)

    main_legends = _extract_main_figure_legends(figure_legends_md_path.read_text(encoding="utf-8"))
    table1 = _demote_table_heading(table1_md_path.read_text(encoding="utf-8"))

    out = (
        manuscript.strip()
        + "\n\n"
        + _PAGE_BREAK_OPENXML
        + "\n\n## Figure Legends\n\n"
        + main_legends
        + "\n"
        + _PAGE_BREAK_OPENXML
        + "\n\n## Tables\n\n"
        + table1
        + "\n"
    )
    return out


def _render_markdown_to_docx(*, markdown: str, out_docx: Path) -> None:
    tmp_md = out_docx.with_suffix(".md")
    tmp_md.write_text(markdown, encoding="utf-8")

    try:
        subprocess.run(
            [
                "pandoc",
                "-f",
                "markdown+superscript",
                "-t",
                "docx",
                "-o",
                str(out_docx),
                str(tmp_md),
            ],
            check=True,
        )
    finally:
        tmp_md.unlink(missing_ok=True)


def _format_table_fonts(table, *, font_name: str, font_size_pt: int, bold_header: bool = True) -> None:
    for r_i, row in enumerate(table.rows):
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = font_name
                    run.font.size = Pt(font_size_pt)
                    if bold_header and r_i == 0:
                        run.bold = True


def _add_table_from_csv(doc: Document, csv_path: Path) -> None:
    df = pd.read_csv(csv_path)

    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    for j, col in enumerate(df.columns):
        hdr_cells[j].text = str(col)

    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for j, col in enumerate(df.columns):
            value = row[col]
            row_cells[j].text = "" if pd.isna(value) else str(value)

    _format_table_fonts(table, font_name="Arial", font_size_pt=10, bold_header=True)


def _add_caption(doc: Document, text: str, *, font_name: str = "Arial", font_size_pt: int = 12) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.name = font_name
    run.font.size = Pt(font_size_pt)


def _add_paragraph(doc: Document, text: str, *, font_name: str, font_size_pt: int, italic: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size_pt)
    run.italic = italic


def _parse_efigure_legends(figure_legends_md: str) -> dict[str, str]:
    """
    Returns mapping like {"eFigure 1": "<legend text>", ...}.
    """
    out: dict[str, list[str]] = {}
    current: str | None = None
    for line in figure_legends_md.splitlines():
        if line.startswith("## eFigure "):
            current = line.removeprefix("## ").split(".", 1)[0].strip()
            out[current] = []
            continue
        if current is None:
            continue
        out[current].append(line)

    return {k: "\n".join(v).strip() for k, v in out.items()}


def _build_paperb_supplement_docx(
    *,
    out_docx: Path,
    paper_title: str,
    supplement_outline_md_path: Path,
    figure_legends_md_path: Path,
    supplement_tables_dir: Path,
    supplement_figures_dir: Path,
) -> None:
    doc = Document()
    _set_document_style(doc, font_name="Times New Roman", font_size_pt=10, line_spacing=WD_LINE_SPACING.SINGLE)
    _add_page_numbers(doc)

    _add_caption(doc, "Supplement", font_name="Times New Roman", font_size_pt=14)
    _add_paragraph(doc, f"for: {paper_title}", font_name="Times New Roman", font_size_pt=10)
    _add_paragraph(
        doc,
        "This document contains eMethods, eTables, and eFigures referenced in the main manuscript.",
        font_name="Times New Roman",
        font_size_pt=10,
    )

    doc.add_paragraph()
    _add_caption(doc, "Contents", font_name="Times New Roman", font_size_pt=12)

    # Parse supplement outline for TOC lines.
    toc_lines: list[str] = []
    for raw in supplement_outline_md_path.read_text(encoding="utf-8").splitlines():
        if raw.startswith("- eMethods "):
            toc_lines.append(raw.removeprefix("- ").strip())
        if raw.startswith("- eTable "):
            toc_lines.append(raw.removeprefix("- ").split(":")[0].strip())
        if raw.startswith("- eFigure "):
            toc_lines.append(raw.removeprefix("- ").split(":")[0].strip())

    for line in toc_lines:
        _add_paragraph(doc, line, font_name="Times New Roman", font_size_pt=10)

    doc.add_page_break()

    _add_caption(doc, "eMethods", font_name="Times New Roman", font_size_pt=12)

    e_methods_blocks = [
        (
            "eMethods 1. Evidence core tables and cohort construction",
            "Triads were PET-anchored and paired within prespecified symmetric windows. One triad per participant was selected deterministically to avoid repeated measures, prioritizing availability of the primary cerebrospinal fluid β-amyloid 42/40 definition and minimizing pairwise date gaps.",
        ),
        (
            "eMethods 2. Latent class model (priors, label constraint, computation)",
            "We fit 2-class Bayesian latent class models with a latent amyloid state A*∈{0,1}. In the primary conditional-independence model, positron emission tomography and cerebrospinal fluid statuses were modeled as independent Bernoulli sensors conditional on A*, and plasma strata were modeled as a 3-level categorical distribution conditional on A*. Priors were Beta(1,1) for prevalence and sensor parameters and Dirichlet(1,1,1) for plasma-category probabilities. In the prespecified PET–cerebrospinal fluid conditional-dependence sensitivity, the joint PET–cerebrospinal fluid 2×2 cell probabilities within each latent class were modeled with Dirichlet(1,1,1,1) priors. Label switching was prevented by enforcing that the A*=1 class has a higher marginal positron emission tomography positivity rate than A*=0 at each iteration. Posterior sampling used Gibbs updates with multiple chains and fixed seeds as specified in the analysis configuration file (outputs/paperB/definitions.yaml).",
        ),
        (
            "eMethods 3. Timing-stratified fits and posterior predictive checks",
            "Timing strata were defined on absolute day gaps (≤7, 8–30, >30 days) and applied separately to each pairwise modality gap. For timing figures, the primary model was refit within each stratum to generate posterior predictive summaries of discordance and mismatch. Observed proportions are shown with 95% Wilson confidence intervals; model summaries use posterior predictive credible intervals. Mismatch was computed among determinate plasma strata only (Low vs High).",
        ),
        (
            "eMethods 4. Prespecified sensitivity analyses (dependence, CSF cutpoint band, plasma discretization)",
            "Sensitivity axes included PET–cerebrospinal fluid conditional dependence, a multiplicative cerebrospinal fluid cutpoint transportability band (0.90× to 1.10×), and an alternative plasma tri-category definition induced by workflow-aligned triage thresholds (as a sensitivity) compared with the primary within-cohort quantile discretization.",
        ),
        (
            "eMethods 5. Coverage/transportability analysis (CSF β-amyloid 42-only)",
            "A prespecified coverage analysis replaced the primary cerebrospinal fluid β-amyloid 42/40 definition with cerebrospinal fluid β-amyloid 42 alone to reduce selection due to incomplete β-amyloid 42/40 availability. Coverage outputs include pattern posteriors and timing-stratified discordance summaries for the full triad cohort.",
        ),
    ]

    for heading, paragraph in e_methods_blocks:
        _add_caption(doc, heading, font_name="Times New Roman", font_size_pt=11)
        _add_paragraph(doc, paragraph, font_name="Times New Roman", font_size_pt=10)

    # eTables
    for i in range(1, 13):
        matches = sorted(supplement_tables_dir.glob(f"eTable{i}_*.csv"))
        if not matches:
            raise FileNotFoundError(f"Missing eTable{i} CSV in {supplement_tables_dir}")
        csv_path = matches[0]
        title = next((t for t in toc_lines if t.startswith(f"eTable {i}.")), f"eTable {i}.")

        doc.add_page_break()
        _add_caption(doc, title, font_name="Arial", font_size_pt=12)
        _add_table_from_csv(doc, csv_path)

    # eFigures
    efigure_legends = _parse_efigure_legends(figure_legends_md_path.read_text(encoding="utf-8"))
    for i in range(1, 9):
        matches = sorted(supplement_figures_dir.glob(f"eFigure{i}_*.png"))
        if not matches:
            raise FileNotFoundError(f"Missing eFigure{i} PNG in {supplement_figures_dir}")
        img = matches[0]
        title = next((t for t in toc_lines if t.startswith(f"eFigure {i}.")), f"eFigure {i}.")
        legend_key = f"eFigure {i}"
        legend_text = efigure_legends.get(legend_key, "")

        doc.add_page_break()
        _add_caption(doc, title, font_name="Arial", font_size_pt=12)
        if legend_text:
            for para in [p.strip() for p in legend_text.split("\n\n") if p.strip()]:
                _add_paragraph(doc, para.replace("\n", " "), font_name="Times New Roman", font_size_pt=9)
        doc.add_picture(str(img), width=Inches(6.7))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    out_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_docx))


def _postprocess_main_docx(docx_path: Path) -> None:
    doc = Document(str(docx_path))
    _set_document_style(doc, font_name="Times New Roman", font_size_pt=12, line_spacing=WD_LINE_SPACING.DOUBLE)
    _add_page_numbers(doc)
    doc.save(str(docx_path))


def main() -> None:
    parser = argparse.ArgumentParser(description="Build Paper B main manuscript + supplement .docx for JAMA submission.")
    parser.add_argument(
        "--manuscript-md",
        type=Path,
        default=Path("manuscript/paperB_neurology/02_manuscript_draft_with_references.md"),
    )
    parser.add_argument(
        "--figure-legends-md",
        type=Path,
        default=Path("manuscript/paperB_neurology/03_figure_legends.md"),
    )
    parser.add_argument(
        "--table1-md",
        type=Path,
        default=Path("outputs/paperB/submission_jama_neurology/main_tables/Table1_Representativeness_IncludedVsExcluded.md"),
    )
    parser.add_argument(
        "--supplement-outline-md",
        type=Path,
        default=Path("outputs/paperB/submission_jama_neurology/SUPPLEMENT_OUTLINE.md"),
    )
    parser.add_argument(
        "--supplement-tables-dir",
        type=Path,
        default=Path("outputs/paperB/submission_jama_neurology/supplement/tables"),
    )
    parser.add_argument(
        "--supplement-figures-dir",
        type=Path,
        default=Path("outputs/paperB/submission_jama_neurology/supplement/figures"),
    )
    parser.add_argument(
        "--out-main-docx",
        type=Path,
        default=Path("outputs/paperB/submission_jama_neurology/PaperB_Manuscript.docx"),
    )
    parser.add_argument(
        "--out-supplement-docx",
        type=Path,
        default=Path("outputs/paperB/submission_jama_neurology/PaperB_Supplement.docx"),
    )
    parser.add_argument(
        "--paper-title",
        type=str,
        default="Amyloid Probability in Alzheimer Disease From Plasma, Cerebrospinal Fluid, and Amyloid Imaging",
    )
    args = parser.parse_args()

    main_md = _build_main_markdown(
        manuscript_md_path=args.manuscript_md,
        figure_legends_md_path=args.figure_legends_md,
        table1_md_path=args.table1_md,
    )
    args.out_main_docx.parent.mkdir(parents=True, exist_ok=True)
    _render_markdown_to_docx(markdown=main_md, out_docx=args.out_main_docx)
    _postprocess_main_docx(args.out_main_docx)

    _build_paperb_supplement_docx(
        out_docx=args.out_supplement_docx,
        paper_title=args.paper_title,
        supplement_outline_md_path=args.supplement_outline_md,
        figure_legends_md_path=args.figure_legends_md,
        supplement_tables_dir=args.supplement_tables_dir,
        supplement_figures_dir=args.supplement_figures_dir,
    )


if __name__ == "__main__":
    main()
