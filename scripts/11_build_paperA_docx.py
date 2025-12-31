from __future__ import annotations

import argparse
import re
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches


def _add_markdown(doc: Document, markdown: str) -> None:
    """
    Minimal Markdown → docx renderer supporting:
    - #/##/### headings
    - bullet lists (- )
    - numbered lists (1. )
    - paragraphs (blank-line separated)
    """

    lines = markdown.splitlines()

    def flush_paragraph(buffer: list[str]) -> None:
        if not buffer:
            return
        text = " ".join([ln.strip() for ln in buffer if ln.strip()])
        if text:
            doc.add_paragraph(text)
        buffer.clear()

    buf: list[str] = []
    for raw in lines:
        line = raw.rstrip("\n")
        if not line.strip():
            flush_paragraph(buf)
            continue

        if line.startswith("# "):
            flush_paragraph(buf)
            doc.add_heading(line[2:].strip(), level=1)
            continue
        if line.startswith("## "):
            flush_paragraph(buf)
            doc.add_heading(line[3:].strip(), level=2)
            continue
        if line.startswith("### "):
            flush_paragraph(buf)
            doc.add_heading(line[4:].strip(), level=3)
            continue

        if line.startswith("- "):
            flush_paragraph(buf)
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
            continue

        if re.match(r"^\\d+\\.\\s+", line):
            flush_paragraph(buf)
            doc.add_paragraph(re.sub(r"^\\d+\\.\\s+", "", line).strip(), style="List Number")
            continue

        # Strip simple bold markers to avoid raw markdown in the doc.
        line = line.replace("**", "")
        buf.append(line)

    flush_paragraph(buf)


def _add_table_from_csv(doc: Document, csv_path: Path, caption: str) -> None:
    df = pd.read_csv(csv_path)
    doc.add_paragraph(caption, style="Intense Quote")

    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    for j, col in enumerate(df.columns):
        hdr_cells[j].text = str(col)

    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for j, col in enumerate(df.columns):
            value = row[col]
            row_cells[j].text = "" if pd.isna(value) else str(value)


def _add_figure(doc: Document, image_path: Path, legend: str, *, width_in: float = 6.5) -> None:
    doc.add_paragraph(legend, style="Intense Quote")
    if image_path.exists():
        doc.add_picture(str(image_path), width=Inches(width_in))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def main() -> None:
    parser = argparse.ArgumentParser(description="Build a Paper A Neurology draft .docx from manuscript assets.")
    parser.add_argument(
        "--manuscript-md",
        type=Path,
        default=Path("manuscript/paperA_neurology/paperA_full_manuscript.md"),
    )
    parser.add_argument(
        "--out-docx",
        type=Path,
        default=Path("manuscript/paperA_neurology/paperA_neurology_draft.docx"),
    )
    parser.add_argument(
        "--tables-dir",
        type=Path,
        default=Path("manuscript/paperA_neurology/tables"),
    )
    parser.add_argument(
        "--figures-dir",
        type=Path,
        default=Path("outputs/paperA/figures"),
    )
    args = parser.parse_args()

    manuscript_md: Path = args.manuscript_md
    out_docx: Path = args.out_docx
    tables_dir: Path = args.tables_dir
    figures_dir: Path = args.figures_dir

    doc = Document()

    md_text = manuscript_md.read_text(encoding="utf-8")
    _add_markdown(doc, md_text)

    # Tables (main + key supplement)
    doc.add_page_break()
    doc.add_heading("Tables", level=1)

    _add_table_from_csv(
        doc,
        tables_dir / "table1_cohort_characteristics.csv",
        "Table 1. Cohort characteristics (evaluation set; one triad per participant).",
    )
    doc.add_paragraph((tables_dir / "table1_footnote.txt").read_text(encoding="utf-8").strip())

    doc.add_paragraph("")
    _add_table_from_csv(
        doc,
        tables_dir / "table2_primary_metrics.csv",
        "Table 2. Primary agreement and workflow metrics (op_ppa0.95_npa0.95; indeterminate→negative).",
    )
    doc.add_paragraph((tables_dir / "table2_footnote.txt").read_text(encoding="utf-8").strip())

    doc.add_paragraph("")
    _add_table_from_csv(
        doc,
        tables_dir / "table3_benchmark_swap_evaluation_primary.csv",
        "Table 3. Primary paired benchmark swap (evaluation-only; CSF-A minus PET; op_ppa0.95_npa0.95; indeterminate→negative).",
    )

    doc.add_paragraph("")
    _add_table_from_csv(
        doc,
        tables_dir / "tableS2_benchmark_swap_pooled_primary.csv",
        "Table S2. Precision-enhancing pooled paired benchmark swap (locked thresholds; CSF-A minus PET; op_ppa0.95_npa0.95; indeterminate→negative).",
    )

    # Figures (embed PNGs for easy review)
    doc.add_page_break()
    doc.add_heading("Figures", level=1)

    legends_path = Path("manuscript/paperA_neurology/03_figure_legends.md")
    legends_text = legends_path.read_text(encoding="utf-8")
    _add_markdown(doc, legends_text)

    doc.add_paragraph("")
    _add_figure(
        doc,
        figures_dir / "figure1_cohort_flow.png",
        "Figure 1. STARD-style cohort flow (embedded preview).",
    )

    doc.add_paragraph("")
    _add_figure(
        doc,
        figures_dir / "figure2_benchmark_swap_forest_evaluation_one_triad_per_rid.png",
        "Figure 2. Primary benchmark swap forest plot (evaluation-only; embedded preview).",
    )

    doc.add_paragraph("")
    _add_figure(
        doc,
        figures_dir / "figure2_benchmark_swap_forest_pooled_one_triad_per_rid.png",
        "Figure S2. Pooled swap-eligible forest plot (locked thresholds; embedded preview).",
    )

    out_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_docx))


if __name__ == "__main__":
    main()

